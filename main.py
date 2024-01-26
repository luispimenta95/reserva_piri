from docx import Document as doc
import os
from dotenv import load_dotenv
import funcoes
from datetime import datetime
from spire.doc import *
from spire.doc.common import *
# Carrega as variáveis de ambiente do arquivo .env
load_dotenv()

today = datetime.today()
months = ['Janeiro', 'Fevereiro', 'Março', 'Abril', 'Maio', 'Junho', 'Julho', 'Agosto', 'Setembro', 'Outubro',
          'Novembro', 'Dezembro']

str_month = months[
    today.month - 1]  # obtemos o numero do mês e subtraímos 1 para que haja a correspondência correta com a nossa lista de meses

cidade = os.getenv('cidade')
estado = os.getenv('estado')
dono = os.getenv('nome_propietario')
cpf = os.getenv('cpf')
bloco = os.getenv('boloco')
email = os.getenv('email')
telefone = os.getenv('telefone')
fracao = os.getenv('fracao')
tipo_ap = os.getenv('tipo')
num_ap = os.getenv('num_ap')


record = int(input("Informe o numero de hospédes (No maximo 4) :"))
if record > 4:
    print("São permitidos somente 4 hospédes.")
    record = 4

pessoa = dict()
pessoas = list()

for i in range(0,record):
    pessoa['nome'] = input(f"Digite o nome da {i+1}ª pessoa :")
    pessoa['cpf'] = input(f"Digite o cpf da {i+1}ª pessoa :")
    pessoa['nascimento'] = input(f"Digite a data de nascimento da {i+1}ª pessoa :")

    pessoas.append(pessoa.copy())
    dataIni = input("Informe a data do check-in :")
    dataFim = input("Informe a data do check-out :")

    referencias = {
        "NOME_USUARIO": dono,
        "CIDADE": cidade,
        "ESTADO": estado,
        "DD": str(datetime.now().day),
        "MM": str_month,
        "AAAA": str(datetime.now().year),
        "CPF_USUARIO": cpf,
        "EMAIL_USUARIO": email,
        "TELEFONE_USUARIO": telefone,
        "END_BLOCO": bloco,
        "DATA_ENTRADA" : dataIni,
        "DATA_SAIDA": dataFim,
        "TIPO_AP" : tipo_ap,
        "NUM_AP" : num_ap,
        "TIPO_FRACAO": fracao


    }
documento = doc('contratos/modelo_reserva.docx')
tabelaDados = documento.tables[0]
tabelaHospedes = documento.tables[1]
last_row = tabelaHospedes.rows[-1]

for row in tabelaDados.rows:
    for cell in row.cells:
        for codigo in referencias:
            valor = referencias[codigo]
            cell.text = cell.text.replace(codigo, valor)

for dado in pessoas:
    cpf = funcoes.formataCpf(dado['cpf'])
    # Adding a row and then adding data in it.
    row = tabelaHospedes.add_row().cells
    # Converting id to string as table can only take string input
    row[0].text = dado['nome']
    row[1].text = cpf
    row[2].text = dado['nascimento']


for paragrafo in documento.paragraphs:
    for codigo in referencias:
        valor = referencias[codigo]
        paragrafo.text = paragrafo.text.replace(codigo,valor)
nomeArquivo = f"{pessoas[0]['nome']}_{dataIni.replace('/','_')}.docx"
documento.save(nomeArquivo)
funcoes.montarPdf(nomeArquivo)